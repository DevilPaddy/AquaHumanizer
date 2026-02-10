"""
AI Humanizer API - Simple Rule-Based Text Enhancement
Copyright (c) 2026

Simple, fast text humanization using regex-based transformations.
No ML models, no complex pipelines - just pure Python stdlib + FastAPI.
"""

import os
import logging
import time
import tempfile
from typing import Optional
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from humanizer import humanize, compute_scores
from format_preserver import extract_text_from_docx, format_output_text

# --- LOGGING CONFIGURATION ---

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- APP INITIALIZATION ---

app = FastAPI(title="AI Humanizer API", version="2.0")

# ── CORS ──────────────────────────────────────────────────────
# allow_credentials MUST be False when allow_origins=["*"]
# Browser will block the request otherwise (CORS spec violation)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,       # ← IMPORTANT: False, not True
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- API MODELS ---

class HumanizeRequest(BaseModel):
    text: str
    style: Optional[str] = "neutral"  # Accept but ignore for now

class HumanizeResponse(BaseModel):
    output:             str  # Changed from humanized_text to match frontend
    human_score:        int
    ai_score:           int
    uniqueness:         int
    word_count_original: int
    word_count_output:  int
    processing_time_ms: int

# --- ENDPOINTS ---

@app.get("/health")
async def health():
    """Health check endpoint"""
    return {"status": "ok", "version": "2.0"}

@app.get("/")
async def root():
    return {"message": "AI Humanizer API is running", "endpoint": "POST /humanize"}

@app.post("/humanize", response_model=HumanizeResponse)
async def humanize_endpoint(req: HumanizeRequest):
    """
    Main endpoint for raw text humanization.
    Uses simple rule-based transformations.
    """
    try:
        # Validate
        text = req.text.strip() if req.text else ""
        if not text:
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        
        word_count = len(text.split())
        if word_count > 8000:
            raise HTTPException(status_code=400, detail="Text too long. Max 8000 words.")
        
        # Process
        start = time.time()
        humanized = humanize(text)
        humanized = format_output_text(humanized)
        scores = compute_scores(text, humanized)
        elapsed_ms = int((time.time() - start) * 1000)
        
        logger.info(f"Humanized {word_count} words → {len(humanized.split())} words in {elapsed_ms}ms")
        
        return HumanizeResponse(
            output=humanized,  # Changed from humanized_text to match frontend
            word_count_original=word_count,
            word_count_output=len(humanized.split()),
            processing_time_ms=elapsed_ms,
            **scores,
        )
    
    except HTTPException:
        # Re-raise HTTP exceptions (validation errors)
        raise
    except Exception as e:
        # Catch all other exceptions, log with full traceback, and return 500
        logger.error(f"Humanization processing error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

@app.post("/humanize-document")
async def humanize_document(
    file: UploadFile = File(..., description="Document file (TXT or DOCX)"),
    style: str = Form("neutral"),
    output_format: str = Form("json", description="Output format: json, txt, or docx")
):
    """
    Handles file uploads and processes documents.
    Supports TXT and DOCX input, returns JSON or file download.
    """
    try:
        # Read file content
        content = await file.read()
        
        # Determine file type and extract text
        filename = file.filename.lower()
        if filename.endswith('.txt'):
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError as e:
                logger.error(f"Failed to decode TXT file: {str(e)}")
                raise HTTPException(
                    status_code=400,
                    detail="Invalid TXT file encoding. Please ensure the file is UTF-8 encoded."
                )
        elif filename.endswith('.docx'):
            # For DOCX, we need python-docx library
            # For now, return error if not available
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(content))
                text = extract_text_from_docx(doc)
            except ImportError:
                raise HTTPException(
                    status_code=400,
                    detail="DOCX support not available. Please upload TXT files or install python-docx."
                )
            except Exception as e:
                logger.error(f"Failed to parse DOCX file: {str(e)}")
                raise HTTPException(
                    status_code=400,
                    detail="Invalid or corrupted DOCX file. Please ensure the file is a valid DOCX document."
                )
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .txt or .docx")
        
        # Validate text
        if not text.strip():
            raise HTTPException(status_code=400, detail="Document is empty")
        
        word_count = len(text.split())
        if word_count > 8000:
            raise HTTPException(status_code=400, detail="Document too long. Max 8000 words.")
        
        # Process
        start = time.time()
        humanized = humanize(text)
        humanized = format_output_text(humanized)
        elapsed_ms = int((time.time() - start) * 1000)
        
        logger.info(f"Humanized document: {word_count} words in {elapsed_ms}ms")
        
        # Return based on requested format
        if output_format == "json":
            return {
                "output": humanized,
                "metrics": {
                    "word_count_original": word_count,
                    "word_count_output": len(humanized.split()),
                    "processing_time_ms": elapsed_ms
                }
            }
        
        elif output_format == "txt":
            return Response(
                content=humanized.encode('utf-8'),
                media_type="text/plain",
                headers={"Content-Disposition": "attachment; filename=humanized.txt"}
            )
        
        elif output_format == "docx":
            # For DOCX output, we need python-docx
            try:
                from docx import Document
                import io
                
                doc = Document()
                for paragraph in humanized.split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
                
                # Save to bytes
                docx_io = io.BytesIO()
                doc.save(docx_io)
                docx_io.seek(0)
                
                return Response(
                    content=docx_io.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": "attachment; filename=humanized.docx"}
                )
            except ImportError:
                raise HTTPException(
                    status_code=400,
                    detail="DOCX output not available. Please use 'json' or 'txt' format."
                )
            except Exception as e:
                logger.error(f"Failed to create DOCX output: {str(e)}")
                raise HTTPException(
                    status_code=500,
                    detail="Failed to generate DOCX output. Please try 'json' or 'txt' format."
                )
        else:
            raise HTTPException(status_code=400, detail="Invalid output format. Use: json, txt, or docx")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document processing error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

@app.post("/humanize-stream")
async def humanize_stream(req: HumanizeRequest):
    """
    Streaming endpoint (legacy compatibility).
    Returns same as /humanize endpoint.
    """
    return await humanize_endpoint(req)

# ── DO NOT add @app.options("/humanize") ──────────────────────
# CORSMiddleware handles OPTIONS preflight automatically.
# A manual OPTIONS handler conflicts with the middleware.

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 7860))
    uvicorn.run(app, host="0.0.0.0", port=port)
